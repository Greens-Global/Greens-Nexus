"""Documents (DMS) - Phase 4: TipTap JSON -> renderable blocks -> PDF/DOCX.
Phase 7: alignment/color/font/hyperlink formatting added to the block model
so those editor features aren't silently stripped on export.

Single AST walk (tiptap_to_blocks) shared by both exporters so table/list/image/
merge-field handling isn't duplicated between the reportlab and python-docx
paths. Unknown TipTap node types (blockquote, codeBlock, horizontalRule - all
StarterKit defaults not explicitly special-cased) fall back to recursing into
`.content` so a stray node type can't crash export.

PDF rendering mirrors routers/esign.py's _build_template_pdf(): reportlab
Platypus flowables (not manual canvas.drawString), LETTER page size, mm-based
margins, and the same _pesc()-style XML-escaping (duplicated for the same
reason merge_fields.py duplicates _merge_data() - a small, self-contained,
compliance-adjacent helper isn't worth an esign.py refactor to reuse).

Known, deliberate approximations (see the plan, not silent gaps):
- PDF font family maps to reportlab's built-in base-14 fonts (no TTF embedding).
- DOCX hyperlinks render styled (colored+underlined) but are not clickable -
  python-docx's public API has no support for the OOXML hyperlink relationship.
"""
import io
from xml.sax.saxutils import escape as _esc
import httpx

from app_url import app_url


# ── Page Setup (Phase 14) ────────────────────────────────────────────────────
# `content.pageSetup` (a sibling of body/header/footer in the same JSON blob -
# no schema change needed) drives both exporters: {size, orientation, margins}.
# Sizes in inches (w, h) as portrait; render_pdf/render_docx swap them for
# landscape. Tabloid has no reportlab.lib.pagesizes constant, so all four are
# defined here directly rather than mixing a library constant with hand-rolled
# ones.
PAGE_SIZES_IN = {
    "letter": (8.5, 11.0),     # US default
    "legal": (8.5, 14.0),      # legal agreements/contracts
    "tabloid": (11.0, 17.0),   # large reports/drawings
    "a4": (8.27, 11.69),       # international
}
# Symmetric margins on all four sides - a simplification of Word's own
# presets (whose "Wide" is asymmetric, 2in sides/1in top-bottom); kept simple
# since both export paths apply one margin value to every side.
MARGINS_IN = {"normal": 1.0, "narrow": 0.5, "wide": 1.5}

DEFAULT_PAGE_SETUP = {"size": "letter", "orientation": "portrait", "margins": "normal"}


def _resolve_page_setup(page_setup: dict = None):
    """(width_in, height_in, margin_in) for the given {size, orientation,
    margins} - unknown/missing values fall back to the US Letter default."""
    ps = {**DEFAULT_PAGE_SETUP, **(page_setup or {})}
    w, h = PAGE_SIZES_IN.get(ps.get("size"), PAGE_SIZES_IN["letter"])
    if ps.get("orientation") == "landscape":
        w, h = h, w
    margin_in = MARGINS_IN.get(ps.get("margins"), MARGINS_IN["normal"])
    return w, h, margin_in


# ── TipTap JSON -> neutral blocks ────────────────────────────────────────────

_EMPTY_RUN_STYLE = {"color": None, "fontFamily": None, "fontSize": None, "link": None,
                     "strike": False, "subscript": False, "superscript": False, "highlight": None}


def _text_runs(nodes, merge: dict) -> list:
    runs = []
    for n in nodes or []:
        t = n.get("type")
        if t == "text":
            marks = n.get("marks") or []
            mark_types = {m.get("type") for m in marks}
            style_attrs = next((m.get("attrs") or {} for m in marks if m.get("type") == "textStyle"), {})
            link_attrs = next((m.get("attrs") or {} for m in marks if m.get("type") == "link"), {})
            highlight_attrs = next((m.get("attrs") or {} for m in marks if m.get("type") == "highlight"), {})
            runs.append({"text": n.get("text") or "", "bold": "bold" in mark_types,
                         "italic": "italic" in mark_types, "underline": "underline" in mark_types,
                         "strike": "strike" in mark_types, "subscript": "subscript" in mark_types,
                         "superscript": "superscript" in mark_types, "highlight": highlight_attrs.get("color"),
                         "color": style_attrs.get("color"), "fontFamily": style_attrs.get("fontFamily"),
                         "fontSize": style_attrs.get("fontSize"), "link": link_attrs.get("href")})
        elif t == "mergeField":
            token = (n.get("attrs") or {}).get("token") or ""
            value = merge.get(token, "{{%s}}" % token)
            runs.append({"text": value, "bold": False, "italic": False, "underline": False, **_EMPTY_RUN_STYLE})
        elif t == "hardBreak":
            runs.append({"text": "\n", "bold": False, "italic": False, "underline": False, **_EMPTY_RUN_STYLE})
        else:
            runs.extend(_text_runs(n.get("content"), merge))
    return runs


def _walk_blocks(nodes, merge: dict) -> list:
    blocks = []
    for n in nodes or []:
        t = n.get("type")
        if t == "paragraph":
            attrs = n.get("attrs") or {}
            align = attrs.get("textAlign") or "left"
            blocks.append({"type": "paragraph", "align": align, "runs": _text_runs(n.get("content"), merge),
                           "border": attrs.get("border"), "spacingBefore": attrs.get("spacingBefore"),
                           "spacingAfter": attrs.get("spacingAfter")})
        elif t == "heading":
            attrs = n.get("attrs") or {}
            blocks.append({"type": "heading", "level": attrs.get("level", 1), "align": attrs.get("textAlign") or "left",
                           "runs": _text_runs(n.get("content"), merge),
                           "border": attrs.get("border"), "spacingBefore": attrs.get("spacingBefore"),
                           "spacingAfter": attrs.get("spacingAfter")})
        elif t in ("bulletList", "orderedList"):
            items = [_walk_blocks(li.get("content"), merge) for li in (n.get("content") or [])]
            blocks.append({"type": t, "items": items})
        elif t == "table":
            rows = []
            for row in n.get("content") or []:
                cells = []
                for cell in (row.get("content") or []):
                    cells.append({
                        "blocks": _walk_blocks(cell.get("content"), merge),
                        "background": (cell.get("attrs") or {}).get("backgroundColor"),
                    })
                rows.append(cells)
            blocks.append({"type": "table", "rows": rows})
        elif t == "image":
            attrs = n.get("attrs") or {}
            src = attrs.get("src") or ""
            if src:
                blocks.append({"type": "image", "src": src, "width": attrs.get("width"), "height": attrs.get("height")})
        elif t == "pageBreak":
            blocks.append({"type": "pageBreak"})
        elif t == "docShape":
            attrs = n.get("attrs") or {}
            blocks.append({"type": "shape", "shapeType": attrs.get("shapeType", "rectangle"),
                           "width": attrs.get("width", 160), "height": attrs.get("height", 100),
                           "fillColor": attrs.get("fillColor", "#dbeafe"), "strokeColor": attrs.get("strokeColor", "#2563eb")})
        elif t == "docTextbox":
            # Floating position/wrap (x/y/wrapMode) is a live-editor-only
            # concept - export keeps textboxes in document order like any
            # other block (documented Phase 9 scope trim, see the plan) and
            # renders the nested content as a bordered single-cell table,
            # reusing the existing table flowable rather than a new primitive.
            blocks.append({"type": "textbox", "blocks": _walk_blocks(n.get("content"), merge)})
        elif t == "mergeField":
            blocks.append({"type": "paragraph", "align": "left", "runs": _text_runs([n], merge)})
        elif t == "sectionBox":
            # The live editor renders this as a bordered box (no export
            # primitive for that yet) - fall back to a bold label line so the
            # section heading isn't silently lost from the PDF/DOCX.
            label = (n.get("attrs") or {}).get("label") or ""
            if label:
                blocks.append({"type": "heading", "level": 3, "align": "left",
                               "runs": [{"text": label, "bold": True, "italic": False, "underline": False, **_EMPTY_RUN_STYLE}]})
            blocks.extend(_walk_blocks(n.get("content"), merge))
        else:
            nested = n.get("content")
            if nested:
                blocks.extend(_walk_blocks(nested, merge))
    return blocks


def tiptap_to_blocks(tiptap_doc, merge: dict) -> list:
    if not tiptap_doc or not isinstance(tiptap_doc, dict):
        return []
    return _walk_blocks(tiptap_doc.get("content") or [], merge)


def _blocks_to_plaintext(blocks: list) -> str:
    """Single-line collapse for header/footer - not full wrapping, just a
    readable one-liner per page (documented tradeoff, see the plan)."""
    lines = []
    for b in blocks or []:
        if b["type"] in ("paragraph", "heading"):
            text = "".join(r["text"] for r in b.get("runs", []))
            if text.strip():
                lines.append(text.strip())
    return "   |   ".join(lines)


def _fetch_image_bytes(url: str):
    if not url:
        return None
    # Letterhead logos may be stored as a path relative to the frontend's own
    # static assets (e.g. "/assets/branding/logo.png") - the browser resolves
    # that against the current origin for free, but this export runs on the
    # backend with no origin of its own, so resolve it against the frontend's
    # public URL (same helper email links already use) before fetching.
    if url.startswith("/"):
        url = app_url() + url
    try:
        resp = httpx.get(url, timeout=10)
        resp.raise_for_status()
        return resp.content
    except Exception:
        return None


def _font_size_num(font_size):
    """'18px' -> 18. Treated as points directly (an approximation already
    baked into this export pipeline - see module docstring)."""
    if not font_size:
        return None
    digits = "".join(c for c in str(font_size) if c.isdigit())
    return int(digits) if digits else None


def _letterhead_text(letterhead: dict, key: str) -> str:
    val = (letterhead or {}).get(key)
    return (val or {}).get("text", "") if isinstance(val, dict) else ""


# ── PDF (reportlab Platypus) ─────────────────────────────────────────────────

# Keyword heuristic (Phase 8) rather than an exhaustive per-name table - the
# frontend's font list grew to ~24 names and reportlab can only ever
# approximate to its built-in base-14 fonts anyway (no TTF embedding), so a
# per-name map would just be maintenance overhead for the same three buckets.
_PDF_SERIF_HINTS = ("times", "georgia", "garamond", "cambria", "palatino", "book antiqua", "serif")
_PDF_MONO_HINTS = ("courier", "consolas", "monaco", "lucida console", "mono")


def _pdf_font_face(family: str) -> str:
    f = (family or "").strip().lower()
    if any(h in f for h in _PDF_MONO_HINTS):
        return "Courier"
    if any(h in f for h in _PDF_SERIF_HINTS):
        return "Times-Roman"
    return "Helvetica"


def _pdf_image_flowable(url: str, max_width_pt: float):
    from reportlab.lib.utils import ImageReader
    from reportlab.platypus import Image as RLImage
    data = _fetch_image_bytes(url)
    if not data:
        return None
    try:
        buf = io.BytesIO(data)
        reader = ImageReader(buf)
        iw, ih = reader.getSize()
        if not iw:
            return None
        scale = min(1.0, max_width_pt / iw)
        buf.seek(0)
        return RLImage(buf, width=iw * scale, height=ih * scale)
    except Exception:
        return None


def _runs_to_markup(runs: list) -> str:
    parts = []
    for r in runs or []:
        t = _esc(r.get("text") or "").replace("\n", "<br/>")
        if r.get("bold"):
            t = f"<b>{t}</b>"
        if r.get("italic"):
            t = f"<i>{t}</i>"
        if r.get("underline"):
            t = f"<u>{t}</u>"
        if r.get("strike"):
            t = f"<strike>{t}</strike>"
        if r.get("superscript"):
            t = f"<super>{t}</super>"
        elif r.get("subscript"):
            t = f"<sub>{t}</sub>"
        font_attrs = []
        if r.get("color"):
            font_attrs.append(f"color='{r['color']}'")
        if r.get("highlight"):
            font_attrs.append(f"backColor='{r['highlight']}'")
        size = _font_size_num(r.get("fontSize"))
        if size:
            font_attrs.append(f"size={size}")
        if r.get("fontFamily"):
            font_attrs.append(f"face='{_pdf_font_face(r['fontFamily'])}'")
        if font_attrs:
            t = f"<font {' '.join(font_attrs)}>{t}</font>"
        if r.get("link"):
            t = f'<a href="{_esc(r["link"])}" color="blue">{t}</a>'
        parts.append(t)
    return "".join(parts)


def _para_style_for_runs(base_style, runs):
    """A run's inline font-size (from a textStyle mark - e.g. an imported
    Word heading with a real 26pt custom font) is applied via `_runs_to_markup`
    as a `<font size=X>` tag INSIDE the paragraph, but reportlab's Paragraph
    flowable allocates vertical space using the STYLE's own fixed `leading`,
    not the actual rendered glyph size. When an inline size exceeds that fixed
    leading, the text visually overflows into whatever flows next - reportlab
    doesn't reflow line height per inline tag the way a browser does. Widening
    the style's fontSize/leading to match the largest size actually used
    (keeping the style's own leading:fontSize ratio) fixes it for whichever
    run is biggest; smaller runs in the same paragraph still render at their
    own explicit inline size, just with more headroom than they need."""
    sizes = [_font_size_num(r.get("fontSize")) for r in (runs or [])]
    sizes = [s for s in sizes if s]
    max_size = max(sizes) if sizes else None
    if not max_size or max_size <= base_style.fontSize:
        return base_style
    ratio = (base_style.leading / base_style.fontSize) if base_style.fontSize else 1.22
    return base_style.clone(f"{base_style.name}_sz{max_size}", fontSize=max_size, leading=max_size * ratio)


_PDF_ALIGN = None  # populated lazily (import needs reportlab, kept local to render functions)


def _pdf_align_style(base_style, align):
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
    align_map = {"left": TA_LEFT, "center": TA_CENTER, "right": TA_RIGHT, "justify": TA_JUSTIFY}
    ta = align_map.get(align)
    if ta is None or base_style.alignment == ta:
        return base_style
    return base_style.clone(f"{base_style.name}_{align}", alignment=ta)


def _shape_drawing(b: dict):
    """Real vector graphics (Phase 8) - reportlab draws these natively, no
    rasterization needed here (contrast with the DOCX path below, which has
    no vector-drawing API and must rasterize via Pillow)."""
    from reportlab.graphics.shapes import Drawing, Rect, Circle, Line, Polygon
    from reportlab.lib.colors import HexColor
    w, h = b.get("width", 160), b.get("height", 100)
    fill = HexColor(b.get("fillColor") or "#dbeafe")
    stroke = HexColor(b.get("strokeColor") or "#2563eb")
    st = b.get("shapeType")
    d = Drawing(w, h)
    if st == "circle":
        d.add(Circle(w / 2, h / 2, min(w, h) / 2 - 2, fillColor=fill, strokeColor=stroke, strokeWidth=2))
    elif st == "line":
        d.add(Line(2, h / 2, w - 2, h / 2, strokeColor=stroke, strokeWidth=3))
    elif st == "triangle":
        d.add(Polygon([w / 2, h - 2, w - 2, 2, 2, 2], fillColor=fill, strokeColor=stroke, strokeWidth=2))
    elif st == "arrow":
        d.add(Line(2, h / 2, w - 14, h / 2, strokeColor=stroke, strokeWidth=3))
        d.add(Polygon([w - 14, h / 2 - 6, w - 2, h / 2, w - 14, h / 2 + 6], fillColor=stroke, strokeColor=stroke))
    else:
        d.add(Rect(2, 2, w - 4, h - 4, rx=4, fillColor=fill, strokeColor=stroke, strokeWidth=2))
    return d


def _para_spacing_style(style, b: dict):
    """spacingBefore/spacingAfter (points, set via the Layout tab's numeric
    fields) map straight onto ParagraphStyle's own space-before/after, which
    reportlab already measures in points - no unit conversion needed."""
    before, after = b.get("spacingBefore"), b.get("spacingAfter")
    if before is None and after is None:
        return style
    return style.clone(f"{style.name}_sp{before}_{after}",
                        spaceBefore=before if before is not None else style.spaceBefore,
                        spaceAfter=after if after is not None else style.spaceAfter)


def _bordered_para(flowable, content_width: float, border: str):
    """Wraps a single Paragraph in a 1-cell borderless-elsewhere Table, the
    same technique already used for text boxes below - reportlab's Paragraph
    flowable has no native border/box of its own to draw one directly."""
    from reportlab.lib import colors
    from reportlab.platypus import Table, TableStyle
    color = colors.HexColor("#111827")
    for hex_str in (border or "").split():
        if hex_str.startswith("#"):
            color = colors.HexColor(hex_str)
    tbl = Table([[flowable]], colWidths=[content_width])
    tbl.setStyle(TableStyle([
        ("BOX", (0, 0), (-1, -1), 1, color),
        ("LEFTPADDING", (0, 0), (-1, -1), 6), ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 4), ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    return tbl


def _blocks_to_flow(blocks: list, body_style, styles, content_width: float) -> list:
    from reportlab.lib import colors
    from reportlab.lib.units import mm
    from reportlab.platypus import Paragraph, Spacer, Table, TableStyle, PageBreak, ListFlowable, ListItem
    flow = []
    for b in blocks or []:
        t = b["type"]
        if t == "paragraph":
            style = _pdf_align_style(body_style, b.get("align"))
            style = _para_style_for_runs(style, b["runs"])
            style = _para_spacing_style(style, b)
            para = Paragraph(_runs_to_markup(b["runs"]) or "&nbsp;", style)
            flow.append(_bordered_para(para, content_width, b.get("border")) if b.get("border") else para)
        elif t == "heading":
            level = max(1, min(6, b.get("level", 1)))
            style = _pdf_align_style(styles[f"Heading{level}"], b.get("align"))
            style = _para_style_for_runs(style, b["runs"])
            style = _para_spacing_style(style, b)
            para = Paragraph(_runs_to_markup(b["runs"]), style)
            flow.append(_bordered_para(para, content_width, b.get("border")) if b.get("border") else para)
        elif t in ("bulletList", "orderedList"):
            items = []
            for item_blocks in b["items"]:
                sub = _blocks_to_flow(item_blocks, body_style, styles, content_width) or [Paragraph("", body_style)]
                items.append(ListItem(sub))
            flow.append(ListFlowable(items, bulletType="bullet" if t == "bulletList" else "1"))
        elif t == "table":
            rows = b["rows"]
            if rows:
                ncols = max((len(r) for r in rows), default=1)
                data = [[(_blocks_to_flow(cell["blocks"], body_style, styles, content_width) or [Paragraph("", body_style)])
                         for cell in row] for row in rows]
                tbl = Table(data, colWidths=[content_width / max(ncols, 1)] * ncols)
                style_cmds = [
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#d1d5db")),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 6), ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                    ("TOPPADDING", (0, 0), (-1, -1), 4), ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                ]
                # Word-style cell shading (Phase 15) - one BACKGROUND command
                # per shaded cell; unshaded cells are left to the table's
                # default (white) background.
                for ri, row in enumerate(rows):
                    for ci, cell in enumerate(row):
                        if cell.get("background"):
                            style_cmds.append(("BACKGROUND", (ci, ri), (ci, ri), colors.HexColor(cell["background"])))
                tbl.setStyle(TableStyle(style_cmds))
                flow.append(tbl)
                flow.append(Spacer(1, 3 * mm))
        elif t == "image":
            # Word-parity resize (Phase 15) - an explicit width (set by
            # dragging the image's resize handle) is honored in points
            # (px * 72/96, the standard CSS-px-to-point conversion) as a max
            # width instead of always auto-scaling to the column width.
            # Known asymmetry vs. the DOCX path: _pdf_image_flowable always
            # preserves the source image's own aspect ratio and never
            # upscales past its natural size (see that helper) - a
            # non-uniform (stretched) resize in the editor is honored
            # exactly in DOCX but only approximated (width capped, height
            # follows aspect ratio) here.
            if b.get("width"):
                img = _pdf_image_flowable(b["src"], min(b["width"] * 0.75, content_width))
            else:
                img = _pdf_image_flowable(b["src"], content_width)
            if img:
                flow.append(img)
                flow.append(Spacer(1, 3 * mm))
        elif t == "shape":
            flow.append(_shape_drawing(b))
            flow.append(Spacer(1, 3 * mm))
        elif t == "textbox":
            inner = _blocks_to_flow(b.get("blocks"), body_style, styles, content_width * 0.9) or [Paragraph("", body_style)]
            tbl = Table([[inner]], colWidths=[content_width])
            tbl.setStyle(TableStyle([
                ("BOX", (0, 0), (-1, -1), 1, colors.HexColor("#9ca3af")),
                ("LEFTPADDING", (0, 0), (-1, -1), 8), ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ("TOPPADDING", (0, 0), (-1, -1), 6), ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ]))
            flow.append(tbl)
            flow.append(Spacer(1, 3 * mm))
        elif t == "pageBreak":
            flow.append(PageBreak())
    return flow


def _letterhead_flow(letterhead: dict, content_width: float) -> list:
    from reportlab.lib import colors
    from reportlab.lib.units import mm
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.platypus import Paragraph, Spacer, Table, TableStyle, HRFlowable
    if not letterhead:
        return []
    flow = []
    logo_url = letterhead.get("logoPath") or ""
    name = letterhead.get("name") or ""
    address = letterhead.get("address") or ""
    header_text = _letterhead_text(letterhead, "headerJson")
    footer_text = _letterhead_text(letterhead, "footerJson")

    if header_text:
        flow.append(Paragraph(f"<font size=8 color='#6b7280'>{_esc(header_text)}</font>",
                              ParagraphStyle("lh_header", fontName="Helvetica", fontSize=8)))
        flow.append(Spacer(1, 2 * mm))

    text_parts = []
    if name:
        text_parts.append(f"<b>{_esc(name)}</b>")
    if address:
        text_parts.append(f"<font size=9 color='#6b7280'>{_esc(address)}</font>")
    text_flow = Paragraph("<br/>".join(text_parts) or "&nbsp;", ParagraphStyle("lh", fontName="Helvetica", fontSize=13, leading=17))
    logo_img = _pdf_image_flowable(logo_url, 45 * mm) if logo_url else None
    if logo_img:
        t = Table([[logo_img, text_flow]], colWidths=[50 * mm, content_width - 50 * mm])
        t.setStyle(TableStyle([("VALIGN", (0, 0), (-1, -1), "MIDDLE"), ("LEFTPADDING", (0, 0), (-1, -1), 0)]))
        flow.append(t)
    else:
        flow.append(text_flow)

    if footer_text:
        flow.append(Spacer(1, 2 * mm))
        flow.append(Paragraph(f"<font size=8 color='#6b7280'>{_esc(footer_text)}</font>",
                              ParagraphStyle("lh_footer", fontName="Helvetica", fontSize=8)))

    flow.append(Spacer(1, 3 * mm))
    flow.append(HRFlowable(width="100%", thickness=1.3, color=colors.HexColor("#111827")))
    flow.append(Spacer(1, 6 * mm))
    return flow


def render_pdf(title: str, header_blocks: list, pages_blocks: list, footer_blocks: list,
               letterhead: dict = None, page_setup: dict = None) -> bytes:
    from reportlab.lib.units import inch, mm
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak

    w_in, h_in, margin_in = _resolve_page_setup(page_setup)
    pagesize = (w_in * inch, h_in * inch)
    margin = margin_in * inch

    styles = getSampleStyleSheet()
    body_style = ParagraphStyle("body", parent=styles["Normal"], fontName="Helvetica",
                                fontSize=10.5, leading=15, spaceAfter=8)
    content_width = pagesize[0] - 2 * margin

    flow = [Paragraph(_esc(title), ParagraphStyle("title", parent=styles["Title"], fontSize=16)), Spacer(1, 4 * mm)]
    flow.extend(_letterhead_flow(letterhead, content_width))
    # pages_blocks: one block-list per page (Document Builder's Pages panel -
    # each page is now a real, independent unit, not a pageBreak marker
    # inside one continuous body). Force a real page boundary between them;
    # _blocks_to_flow's own pageBreak-block handling still runs too, for
    # backward compat with content saved before pages were real, independent
    # units (a single-page body doc could still carry manual pageBreak nodes).
    for i, blocks in enumerate(pages_blocks):
        if i > 0:
            flow.append(PageBreak())
        flow.extend(_blocks_to_flow(blocks, body_style, styles, content_width))
    if not flow:
        flow.append(Paragraph("", body_style))

    header_text = _blocks_to_plaintext(header_blocks)[:180]
    footer_text = _blocks_to_plaintext(footer_blocks)[:180]

    def decorate(canvas, doc):
        canvas.saveState()
        canvas.setFont("Helvetica", 9)
        canvas.setFillColor(colors.HexColor("#6b7280"))
        if header_text:
            canvas.drawString(margin, pagesize[1] - margin + 8, header_text)
        if footer_text:
            canvas.drawString(margin, margin - 14, footer_text)
        # Right-aligned page number, mirroring a printed SOP's footer (title
        # left, page right). Current-page only (no "of N") - that needs a
        # second pass to know the final page count before the first page is
        # drawn, which SimpleDocTemplate doesn't expose here.
        canvas.drawRightString(pagesize[0] - margin, margin - 14, f"Page {canvas.getPageNumber()}")
        canvas.restoreState()

    buf = io.BytesIO()
    SimpleDocTemplate(buf, pagesize=pagesize, topMargin=margin, bottomMargin=margin,
                      leftMargin=margin, rightMargin=margin, title=title).build(
        flow, onFirstPage=decorate, onLaterPages=decorate)
    return buf.getvalue()


# ── DOCX (python-docx) ────────────────────────────────────────────────────────

_DOCX_ALIGN_NAMES = {"left": "LEFT", "center": "CENTER", "right": "RIGHT", "justify": "JUSTIFY"}


def _docx_align(align):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    name = _DOCX_ALIGN_NAMES.get(align)
    return getattr(WD_ALIGN_PARAGRAPH, name) if name else None


def _hex_to_rgbcolor(hex_color):
    from docx.shared import RGBColor
    try:
        h = (hex_color or "").lstrip("#")
        if len(h) != 6:
            return None
        return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))
    except Exception:
        return None


_DOCX_HIGHLIGHT_MAP = {
    # python-docx only accepts WD_COLOR_INDEX's fixed named palette, not
    # arbitrary hex - each TipTap highlight color snaps to its nearest match.
    "#fef08a": "YELLOW", "#fde68a": "YELLOW", "#fecaca": "RED", "#fca5a5": "RED",
    "#bbf7d0": "BRIGHT_GREEN", "#86efac": "BRIGHT_GREEN", "#bfdbfe": "TURQUOISE",
    "#93c5fd": "TURQUOISE", "#e9d5ff": "VIOLET", "#d8b4fe": "VIOLET",
    "#fed7aa": "DARK_YELLOW", "#fdba74": "DARK_YELLOW",
}


def _docx_highlight(hex_color: str):
    from docx.enum.text import WD_COLOR_INDEX
    name = _DOCX_HIGHLIGHT_MAP.get((hex_color or "").lower(), "YELLOW" if hex_color else None)
    return getattr(WD_COLOR_INDEX, name) if name else None


def _add_runs(paragraph, runs: list) -> None:
    from docx.shared import Pt, RGBColor
    for r in runs or []:
        text = r.get("text") or ""
        if not text:
            continue
        run = paragraph.add_run(text)
        run.bold = bool(r.get("bold"))
        run.italic = bool(r.get("italic"))
        link = r.get("link")
        run.underline = bool(r.get("underline")) or bool(link)  # link: styled, not clickable - see module docstring
        run.font.strike = bool(r.get("strike"))
        run.font.subscript = bool(r.get("subscript"))
        run.font.superscript = bool(r.get("superscript"))
        if r.get("highlight"):
            hl = _docx_highlight(r["highlight"])
            if hl:
                run.font.highlight_color = hl
        color = _hex_to_rgbcolor(r.get("color")) or (RGBColor(0x1a, 0x0d, 0xab) if link and not r.get("color") else None)
        if color:
            run.font.color.rgb = color
        size = _font_size_num(r.get("fontSize"))
        if size:
            run.font.size = Pt(size)
        if r.get("fontFamily"):
            run.font.name = r["fontFamily"]


def _shape_png_bytes(b: dict):
    """Rasterize via Pillow (Phase 8) - python-docx's public API has no
    vector-drawing support (same category of limitation as Phase 7's
    non-clickable DOCX hyperlinks), so this is a deliberate, documented
    approximation rather than a native shape."""
    from PIL import Image, ImageDraw
    w, h = b.get("width", 160), b.get("height", 100)
    img = Image.new("RGBA", (w, h), (0, 0, 0, 0))
    draw = ImageDraw.Draw(img)
    fill = b.get("fillColor") or "#dbeafe"
    stroke = b.get("strokeColor") or "#2563eb"
    st = b.get("shapeType")
    try:
        if st == "circle":
            draw.ellipse([2, 2, w - 2, h - 2], fill=fill, outline=stroke, width=2)
        elif st == "line":
            draw.line([2, h // 2, w - 2, h // 2], fill=stroke, width=3)
        elif st == "triangle":
            draw.polygon([(w / 2, 2), (w - 2, h - 2), (2, h - 2)], fill=fill, outline=stroke)
        elif st == "arrow":
            draw.line([2, h // 2, w - 16, h // 2], fill=stroke, width=3)
            draw.polygon([(w - 16, h // 2 - 6), (w - 2, h // 2), (w - 16, h // 2 + 6)], fill=stroke)
        else:
            draw.rounded_rectangle([2, 2, w - 2, h - 2], radius=6, fill=fill, outline=stroke, width=2)
    except Exception:
        return None
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return buf.getvalue()


def _shade_docx_cell(cell, hex_color: str) -> None:
    """Word-style cell background shading (Phase 15) - python-docx has no
    public API for this (unlike paragraph/run formatting), so it's set via
    the raw <w:shd> OOXML element on the cell's <w:tcPr>, the same technique
    Word itself uses under the hood for "Table > Shading"."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    hex_color = (hex_color or "").lstrip("#")
    if len(hex_color) != 6:
        return
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _border_docx_paragraph(paragraph, border: str) -> None:
    """Paragraph border (Layout tab) - same rationale as _shade_docx_cell
    above: python-docx has no public API for a paragraph border, so it's set
    via the raw <w:pBdr> OOXML element on <w:pPr>."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    color = "111827"
    for hex_str in (border or "").split():
        if hex_str.startswith("#") and len(hex_str) == 7:
            color = hex_str.lstrip("#")
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "6")
        el.set(qn("w:space"), "4")
        el.set(qn("w:color"), color)
        pBdr.append(el)
    pPr.append(pBdr)


def _apply_docx_para_extras(paragraph, b: dict) -> None:
    """spacingBefore/spacingAfter + border (Layout tab, editor-only until
    now) - applied after _add_runs so this stays a thin, self-contained
    post-processing step rather than threading extra params through the
    run-formatting call."""
    from docx.shared import Pt
    before, after = b.get("spacingBefore"), b.get("spacingAfter")
    if before is not None:
        paragraph.paragraph_format.space_before = Pt(before)
    if after is not None:
        paragraph.paragraph_format.space_after = Pt(after)
    if b.get("border"):
        _border_docx_paragraph(paragraph, b["border"])


def _blocks_to_docx(doc, blocks: list) -> None:
    from docx.shared import Inches
    for b in blocks or []:
        t = b["type"]
        if t == "paragraph":
            p = doc.add_paragraph()
            align = _docx_align(b.get("align"))
            if align is not None:
                p.alignment = align
            _add_runs(p, b["runs"])
            _apply_docx_para_extras(p, b)
        elif t == "heading":
            level = max(1, min(6, b.get("level", 1)))
            p = doc.add_heading("", level=level)
            align = _docx_align(b.get("align"))
            if align is not None:
                p.alignment = align
            _add_runs(p, b["runs"])
            _apply_docx_para_extras(p, b)
        elif t in ("bulletList", "orderedList"):
            style = "List Bullet" if t == "bulletList" else "List Number"
            for item_blocks in b["items"]:
                if not item_blocks:
                    continue
                first, rest = item_blocks[0], item_blocks[1:]
                if first["type"] == "paragraph":
                    _add_runs(doc.add_paragraph(style=style), first["runs"])
                    _blocks_to_docx(doc, rest)
                else:
                    _blocks_to_docx(doc, item_blocks)
        elif t == "table":
            rows = b["rows"]
            if rows:
                ncols = max((len(r) for r in rows), default=1)
                table = doc.add_table(rows=len(rows), cols=ncols)
                table.style = "Table Grid"
                for ri, row in enumerate(rows):
                    for ci, cell in enumerate(row):
                        docx_cell = table.cell(ri, ci)
                        if cell.get("background"):
                            _shade_docx_cell(docx_cell, cell["background"])
                        first = True
                        for cb in cell["blocks"]:
                            if cb["type"] != "paragraph":
                                continue
                            p = docx_cell.paragraphs[0] if first else docx_cell.add_paragraph()
                            _add_runs(p, cb["runs"])
                            first = False
                doc.add_paragraph()
        elif t == "image":
            data = _fetch_image_bytes(b["src"])
            if data:
                try:
                    # Word-parity resize (Phase 15) - an explicit width/height
                    # (set by dragging the image's resize handle) is honored
                    # exactly (px/96 = inches, same convention `shape` below
                    # already uses). Passing ONLY width to add_picture makes
                    # python-docx auto-scale height to the source image's own
                    # aspect ratio, silently ignoring a non-uniform resize -
                    # both must be passed explicitly when both are known.
                    kwargs = {}
                    if b.get("width"):
                        kwargs["width"] = Inches(b["width"] / 96)
                    if b.get("height"):
                        kwargs["height"] = Inches(b["height"] / 96)
                    if not kwargs:
                        kwargs["width"] = Inches(5.5)
                    doc.add_picture(io.BytesIO(data), **kwargs)
                except Exception:
                    pass
        elif t == "shape":
            data = _shape_png_bytes(b)
            if data:
                try:
                    doc.add_picture(io.BytesIO(data), width=Inches(b.get("width", 160) / 96))
                except Exception:
                    pass
        elif t == "textbox":
            table = doc.add_table(rows=1, cols=1)
            table.style = "Table Grid"
            cell = table.cell(0, 0)
            first = True
            for cb in b.get("blocks") or []:
                if cb["type"] not in ("paragraph", "heading"):
                    continue
                p = cell.paragraphs[0] if first else cell.add_paragraph()
                _add_runs(p, cb.get("runs") or [])
                first = False
            doc.add_paragraph()
        elif t == "pageBreak":
            doc.add_page_break()


def render_docx(title: str, header_blocks: list, pages_blocks: list, footer_blocks: list,
                letterhead: dict = None, page_setup: dict = None) -> bytes:
    from docx import Document as DocxDocument
    from docx.shared import Pt, Inches
    from docx.enum.section import WD_ORIENT

    w_in, h_in, margin_in = _resolve_page_setup(page_setup)

    doc = DocxDocument()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE if (page_setup or {}).get("orientation") == "landscape" else WD_ORIENT.PORTRAIT
    # python-docx's .orientation flag is metadata only - it does not swap
    # page_width/page_height itself, so w_in/h_in (already swapped by
    # _resolve_page_setup for landscape) are what actually reshape the page.
    section.page_width = Inches(w_in)
    section.page_height = Inches(h_in)
    section.left_margin = Inches(margin_in)
    section.right_margin = Inches(margin_in)
    section.top_margin = Inches(margin_in)
    section.bottom_margin = Inches(margin_in)

    header_text = _blocks_to_plaintext(header_blocks)
    if header_text:
        section.header.paragraphs[0].text = header_text
    footer_text = _blocks_to_plaintext(footer_blocks)
    if footer_text:
        section.footer.paragraphs[0].text = footer_text

    doc.add_heading(title, level=1)

    if letterhead:
        lh_header = _letterhead_text(letterhead, "headerJson")
        if lh_header:
            p = doc.add_paragraph(lh_header)
            if p.runs:
                p.runs[0].font.size = Pt(8)
        logo_url = letterhead.get("logoPath") or ""
        if logo_url:
            data = _fetch_image_bytes(logo_url)
            if data:
                try:
                    doc.add_picture(io.BytesIO(data), width=Inches(1.8))
                except Exception:
                    pass
        name = letterhead.get("name") or ""
        if name:
            p = doc.add_paragraph()
            run = p.add_run(name)
            run.bold = True
            run.font.size = Pt(13)
        address = letterhead.get("address") or ""
        if address:
            p = doc.add_paragraph(address)
            if p.runs:
                p.runs[0].font.size = Pt(9)
        lh_footer = _letterhead_text(letterhead, "footerJson")
        if lh_footer:
            p = doc.add_paragraph(lh_footer)
            if p.runs:
                p.runs[0].font.size = Pt(8)
        doc.add_paragraph()

    # pages_blocks: one block-list per page (see render_pdf's matching
    # comment) - a real page break between each one, on top of
    # _blocks_to_docx's own pageBreak-block handling (backward compat).
    for i, blocks in enumerate(pages_blocks):
        if i > 0:
            doc.add_page_break()
        _blocks_to_docx(doc, blocks)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
